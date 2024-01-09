import docx

doc = docx.Document('draft.docx')
#print(doc.paragraphs[0].__sizeof__())
#print(doc.paragraphs.__contains__('a'))
#print(doc.sections[0])
#print(doc.add_page_break())
print(help(doc))
r = 1
#for i in range(27):
#    print(doc.paragraphs[i].text)
for i in range(len(doc.paragraphs)):
    if '1' in doc.paragraphs[i].text:
        print(f'Совпадение {r}:',doc.paragraphs[i].text)
        r = r+1

#print(doc.paragraphs[0].text)